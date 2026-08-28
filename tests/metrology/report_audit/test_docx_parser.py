"""Tests DocxParser against a real, generated .docx fixture."""
from __future__ import annotations

import docx

from ai_tools_lab.domains.metrology.report_audit.ingestion.docx_parser import parse_docx


def test_parse_docx_extracts_paragraphs_and_tables(tmp_path) -> None:
    document = docx.Document()
    document.add_paragraph("报告编号: JL-2")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "标准值"
    table.rows[0].cells[1].text = "示值"
    table.rows[1].cells[0].text = "0"
    table.rows[1].cells[1].text = "0.1"

    path = tmp_path / "sample.docx"
    document.save(path)

    raw_document = parse_docx(path, file_hash="abc")

    assert "报告编号: JL-2" in raw_document.pages[0].text
    assert raw_document.pages[0].tables[0].rows[0] == ["标准值", "示值"]
