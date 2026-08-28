"""Parses Word (.docx) reports into a RawDocument using python-docx."""
from __future__ import annotations

from pathlib import Path

import docx

from ai_tools_lab.domains.metrology.report_audit.models.raw_document import (
    RawDocument,
    RawPage,
    RawTable,
)


def parse_docx(path: str | Path, file_hash: str) -> RawDocument:
    """Extract paragraph text and tables from a .docx file.

    Word documents have no reliable page boundaries in the XML, so the whole
    document is treated as a single logical page (index 0).
    """
    document = docx.Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = [
        RawTable(
            index=table_index,
            page=0,
            rows=[[cell.text for cell in row.cells] for row in table.rows],
        )
        for table_index, table in enumerate(document.tables)
    ]
    page = RawPage(index=0, text=text, tables=tables)
    return RawDocument(source_path=str(path), file_hash=file_hash, pages=[page])
